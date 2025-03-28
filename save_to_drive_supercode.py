import google.generativeai as genai
from pathlib import Path
import streamlit as st
import json
import time
import re
import os
from dotenv import load_dotenv
from PIL import Image
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# Load environment variables from .env file
load_dotenv()

# Configure GenAI API key
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Function to authenticate Google Drive
def authenticate_drive():
    creds = Credentials.from_authorized_user_file('token.json', ['https://www.googleapis.com/auth/drive.file'])
    service = build('drive', 'v3', credentials=creds)
    return service

# Function to save Word file to Google Drive
def save_to_drive(word_file, file_name):
    service = authenticate_drive()
    file_metadata = {'name': file_name}
    media = MediaIoBaseUpload(word_file, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return file.get('id')

# Function to send email with the Word document as an attachment
def send_email(word_file, email_to, subject="Hemodialysis Data Report"):
    email_from = os.getenv('EMAIL_USER')
    email_password = os.getenv('EMAIL_PASSWORD')
    
    msg = MIMEMultipart()
    msg['From'] = email_from
    msg['To'] = email_to
    msg['Subject'] = subject

    body = "Please find attached the hemodialysis data report."
    msg.attach(MIMEText(body, 'plain'))

    attachment = MIMEBase('application', 'octet-stream')
    word_file.seek(0)
    attachment.set_payload(word_file.read())
    encoders.encode_base64(attachment)
    attachment.add_header('Content-Disposition', f'attachment; filename="hemodialysis_report.docx"')
    msg.attach(attachment)

    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(email_from, email_password)
    text = msg.as_string()
    server.sendmail(email_from, email_to, text)
    server.quit()

# Function to initialize the model
def initialize_model():
    generation_config = {
        "temperature": 0.2,
        "top_p": 1,
        "top_k": 1,
        "max_output_tokens": 2048,
    }
    return genai.GenerativeModel("gemini-1.5-pro", generation_config=generation_config)

# Function to clean the JSON response
def clean_json_response(response_text):
    cleaned_text = re.sub(r'^json\s*```|```$', '', response_text, flags=re.IGNORECASE | re.MULTILINE)
    cleaned_text = cleaned_text.strip()
    return cleaned_text

# Function to process the image and extract hemodialysis data
def extract_hemodialysis_data(model, image):
    image_bytes = io.BytesIO()
    image.save(image_bytes, format='JPEG')
    image_part = {
        "mime_type": "image/jpeg",
        "data": image_bytes.getvalue()
    }
    
    prompt = """
    Analyze this image of a hemodialysis machine display. Extract all visible numerical data and parameters.
    Return the data in a JSON format with keys being the parameter names and values being the numerical readings.
    If you're unsure about a value, use null. Ensure your response is valid JSON without any additional text or formatting.
    """
    
    start_time = time.time()
    response = model.generate_content([prompt, image_part])
    end_time = time.time()
    
    execution_time = end_time - start_time
    
    if response.text:
        cleaned_response = clean_json_response(response.text)
        try:
            data = json.loads(cleaned_response)
            return {"success": True, "data": data}, execution_time
        except json.JSONDecodeError:
            return {"success": False, "error": "Failed to parse JSON response", "raw_response": cleaned_response}, execution_time
    else:
        return {"success": False, "error": "No valid content generated"}, execution_time

# Function to create a Word document from the descriptions
def create_word_file(data):
    doc = Document()
    doc.add_heading('Hemodialysis Data Report', 0)

    if isinstance(data, dict):
        for key, value in data.items():
            doc.add_paragraph(f"{key}: {value}")
    else:
        doc.add_paragraph("No valid data available.")

    # Save the Word file to a BytesIO object
    word_file = io.BytesIO()
    doc.save(word_file)
    word_file.seek(0)

    return word_file

# Streamlit app
def main():
    st.title("Hemodialysis Data Extractor")

    # Initialize session state
    if 'history' not in st.session_state:
        st.session_state.history = []
    if 'model' not in st.session_state:
        st.session_state.model = initialize_model()

    # Sidebar for navigation
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["Data Extraction", "History"])

    if page == "Data Extraction":
        st.header("Upload and Extract Data")

        # Upload an image file
        uploaded_file = st.file_uploader("Choose an image of a hemodialysis machine display", type=["jpg", "jpeg", "png"])

        if uploaded_file is not None:
            # Display the uploaded image
            image = Image.open(uploaded_file)
            st.image(image, caption='Uploaded Image', use_column_width=True)
            
            # Button to extract data
            if st.button("Extract Data"):
                result, execution_time = extract_hemodialysis_data(st.session_state.model, image)
                
                # Display the extracted data or error message
                st.write("Extraction Result:")
                if result["success"]:
                    st.json(result["data"])
                    # Create a Word document from the results
                    word_file = create_word_file(result["data"])
                else:
                    st.error(f"Error: {result['error']}")
                    if "raw_response" in result:
                        st.write("Raw response from the model:")
                        st.code(result["raw_response"], language="json")
                    # Create an empty Word document since no valid data was found
                    word_file = create_word_file(None)
                
                # Display execution time
                st.write(f"Execution Time: {execution_time:.2f} seconds")

                # Add to history
                st.session_state.history.append({
                    "image": image,
                    "result": result,
                    "execution_time": execution_time
                })

                st.success("Data extracted and added to history!")

                # Email the Word document
                email_to = st.text_input("Enter email address to send the document:")
                if st.button("Send Email"):
                    if email_to:
                        send_email(word_file, email_to)
                        st.success(f"Email sent to {email_to}")
                    else:
                        st.error("Please enter a valid email address.")

                # Save the Word document to Google Drive
                if st.button("Save to Google Drive"):
                    file_id = save_to_drive(word_file, "hemodialysis_report.docx")
                    if file_id:
                        st.success(f"Document saved to Google Drive with ID: {file_id}")
                    else:
                        st.error("Failed to save document to Google Drive")

    elif page == "History":
        st.header("Extraction History")
        if not st.session_state.history:
            st.write("No extraction history available.")
        else:
            for i, entry in enumerate(st.session_state.history):
                st.subheader(f"Entry {i+1}")
                st.image(entry["image"], caption=f'Image {i+1}', use_column_width=True)
                if entry["result"]["success"]:
                    st.json(entry["result"]["data"])
                else:
                    st.error(f"Error: {entry['result']['error']}")
                    if "raw_response" in entry["result"]:
                        st.write("Raw response from the model:")
                        st.code(entry["result"]["raw_response"], language="json")
                st.write(f"Execution Time: {entry['execution_time']:.2f} seconds")
                st.markdown("---")

if __name__ == "__main__":
    main()