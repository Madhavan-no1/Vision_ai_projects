import google.generativeai as genai
from pathlib import Path
import streamlit as st
import json
import time
import re
import os
from dotenv import load_dotenv
from PIL import Image  
import io
from docx import Document  
from docx.shared import Inches


load_dotenv()


genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))


def initialize_model():
    generation_config = {
        "temperature": 0.2,
        "top_p": 1,
        "top_k": 1,
        "max_output_tokens": 2048,
    }
    return genai.GenerativeModel("gemini-1.5-pro", generation_config=generation_config)


def clean_json_response(response_text):
    cleaned_text = re.sub(r'^json\s*```|```$', '', response_text, flags=re.IGNORECASE | re.MULTILINE)
    cleaned_text = cleaned_text.strip()
    return cleaned_text


def extract_hemodialysis_data(model, image):
    image_bytes = io.BytesIO()
    image.save(image_bytes, format='JPEG')
    image_part = {
        "mime_type": "image/jpeg",
        "data": image_bytes.getvalue()
    }
    
    prompt = """
    Analyze this image of a hemodialysis machine display. Extract all visible numerical data, name, age as of 2024and parameters.
    Return the data in a JSON format with keys being the name in image , parameter names and values being the numerical readings.
    If you're unsure about a value, use null. Ensure your response is valid JSON without any additional text or formatting.
    """
    
    start_time = time.time()
    response = model.generate_content([prompt, image_part])
    end_time = time.time()
    
    execution_time = end_time - start_time
    
    if response.text:
        cleaned_response = clean_json_response(response.text)
        try:
            data = json.loads(cleaned_response)
            return {"success": True, "data": data}, execution_time
        except json.JSONDecodeError:
            return {"success": False, "error": "Failed to parse JSON response", "raw_response": cleaned_response}, execution_time
    else:
        return {"success": False, "error": "No valid content generated"}, execution_time


def create_word_document(data, image):
    document = Document()
    
   
    document.add_heading('Hemodialysis Data Extraction', 0)
    
    #uploaded image
    image_bytes = io.BytesIO()
    image.save(image_bytes, format='JPEG')
    image_bytes.seek(0)
    document.add_picture(image_bytes, width=Inches(6))
    document.add_paragraph('Uploaded Image')
    
    #table for the data
    document.add_heading('Extracted Data', level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'
    
    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    
    # Save the document
    doc_bytes = io.BytesIO()
    document.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

# Streamlit
def main():
    st.set_page_config(page_title="Hemodialysis Data Extractor", layout="wide")
    st.title("Hemodialysis Data Extractor")

    # Initialize
    if 'history' not in st.session_state:
        st.session_state.history = []
    if 'model' not in st.session_state:
        st.session_state.model = initialize_model()

    
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["Data Extraction", "History"])

    if page == "Data Extraction":
        st.header("Upload and Extract Data")

        
        uploaded_file = st.file_uploader("Choose an image of a hemodialysis machine display", type=["jpg", "jpeg", "png"])

        if uploaded_file is not None:
            
            image = Image.open(uploaded_file)
            st.image(image, caption='Uploaded Image', use_column_width=True)
            
            
            if st.button("Extract Data"):
                with st.spinner("Extracting data..."):
                    result, execution_time = extract_hemodialysis_data(st.session_state.model, image)
                
                
                st.write("### Extraction Result:")
                if result["success"]:
                    st.json(result["data"])
                    
                   
                    doc_bytes = create_word_document(result["data"], image)
                    
                  
                    st.download_button(
                        label="Download as Word Document",
                        data=doc_bytes,
                        file_name="hemodialysis_data.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error(f"Error: {result['error']}")
                    if "raw_response" in result:
                        st.write("Raw response from the model:")
                        st.code(result["raw_response"], language="json")
                
               
                st.write(f"**Execution Time:** {execution_time:.2f} seconds")

                if result["success"]:
                    
                    st.session_state.history.append({
                        "image": image,
                        "result": result,
                        "execution_time": execution_time
                    })

                    st.success("Data extracted and added to history!")

    elif page == "History":
        st.header("Extraction History")
        if not st.session_state.history:
            st.write("No extraction history available.")
        else:
            for i, entry in enumerate(st.session_state.history):
                with st.expander(f"Entry {i+1}"):
                    st.image(entry["image"], caption=f'Image {i+1}', use_column_width=True)
                    if entry["result"]["success"]:
                        st.json(entry["result"]["data"])
                        
                       
                        doc_bytes = create_word_document(entry["result"]["data"], entry["image"])
                        
                        
                        st.download_button(
                            label=f"Download Entry {i+1} as Word Document",
                            data=doc_bytes,
                            file_name=f"hemodialysis_data_entry_{i+1}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error(f"Error: {entry['result']['error']}")
                        if "raw_response" in entry["result"]:
                            st.write("Raw response from the model:")
                            st.code(entry["result"]["raw_response"], language="json")
                    st.write(f"**Execution Time:** {entry['execution_time']:.2f} seconds")
    
   
    st.markdown("---")
    st.markdown("© 2024 Hemodialysis Data Extractor App")

if __name__ == "__main__":
    main()
